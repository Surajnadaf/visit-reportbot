"""
FREE FIELD VISIT REPORT BOT (FINAL STRUCTURE FIXED)
- No AI
- Clean MIS format
- Correct heading and order
"""

import os
from datetime import datetime

from telegram import Update
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    ConversationHandler,
    ContextTypes,
    filters,
)

from docx import Document
from docx.shared import Inches

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet


# =========================
# CONFIG
# =========================

TOKEN = "8876576425:AAE7AyCrtm6VitooNezwHHhoFyOwOCn6JuQ"


# =========================
# STATES
# =========================

LOCATION, VISIT_DATE, OBJECTIVE, OBSERVATIONS, ACTIONS, PHOTOS = range(6)


# =========================
# START
# =========================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data.clear()
    await update.message.reply_text("📍 Location of Visit?")
    return LOCATION


async def location(update, context):
    context.user_data["location"] = update.message.text
    await update.message.reply_text("📅 Visit Date? (Example: 20-06-2026)")
    return VISIT_DATE

async def visit_date(update, context):
    context.user_data["visit_date"] = update.message.text
    await update.message.reply_text("🎯 Objective of Visit?")
    return OBJECTIVE


async def objective(update, context):
    context.user_data["objective"] = update.message.text
    await update.message.reply_text("🧠 Key Observations / Discussion Points?")
    return OBSERVATIONS


async def observations(update, context):
    context.user_data["observations"] = update.message.text
    await update.message.reply_text("📌 Action Items?")
    return ACTIONS


async def actions(update, context):
    context.user_data["actions"] = update.message.text
    context.user_data["photos"] = []

    await update.message.reply_text("📸 Upload up to 8 photos. Send /done when finished.")
    return PHOTOS


# =========================
# PHOTO HANDLER
# =========================

async def photo_handler(update, context):
    photos = context.user_data["photos"]

    if len(photos) >= 8:
        await update.message.reply_text("Max 8 photos allowed.")
        return PHOTOS

    photo = update.message.photo[-1]
    file = await photo.get_file()

    folder = os.path.abspath("downloads")
    os.makedirs(folder, exist_ok=True)

    path = os.path.join(folder, f"photo_{len(photos)+1}.jpg")

    await file.download_to_drive(path)

    photos.append(path)

    await update.message.reply_text(f"Photo {len(photos)} saved.")

    return PHOTOS


# =========================
# SIMPLE SUMMARY (NO AI)
# =========================

def generate_summary(data):

    return f"""
A field visit was conducted at {data.get('location','')}. 
The objective of the visit was {data.get('objective','')}.

The visit focused on reviewing field-level implementation and service delivery.

Key observations highlighted operational conditions and implementation status at the field level.

Action points identified during the visit aim to improve service delivery and strengthen implementation effectiveness.
""".strip()


# =========================
# DOCX GENERATION (FINAL FORMAT)
# =========================

def generate_docx(data):

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Field_Visit_Report_{ts}.docx"

    doc = Document()

    # =========================
    # HEADING FIXED
    # =========================

    doc.add_heading("Field Visit Report", 0)

    # =========================
    # 1. TABLE FIRST (AS REQUESTED)
    # =========================

    doc.add_heading("Field Visit Details", level=1)

    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"

    rows = [
        ("Location Visited", data.get("location", "")),
        ("Visit Date", data.get("visit_date", "")),
        ("Objective of Visit", data.get("objective", "")),
        ("Action Items", data.get("actions", "")),
        ("Report Compiled By", "Suraj Nadaf"),
    ]

    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = str(v)

    # =========================
    # 2. SUMMARY
    # =========================

    doc.add_heading("Summary", level=1)
    doc.add_paragraph(generate_summary(data))

    # =========================
    # 3. OBSERVATIONS
    # =========================

    doc.add_heading("Key Observations", level=1)
    doc.add_paragraph(data.get("observations", ""))

    # =========================
    # 4. PHOTOS
    # =========================

    photos = data.get("photos", [])

    if photos:

        doc.add_page_break()
        doc.add_heading("Photographs", level=1)

        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        for i in range(0, len(photos), 2):

            row = table.add_row().cells

            left = photos[i]

            if os.path.exists(left):
                run = row[0].paragraphs[0].add_run()
                run.add_picture(left, width=Inches(2.5))
            else:
                row[0].text = "Image missing"

            if i + 1 < len(photos):
                right = photos[i + 1]

                if os.path.exists(right):
                    run = row[1].paragraphs[0].add_run()
                    run.add_picture(right, width=Inches(2.5))
                else:
                    row[1].text = "Image missing"

    doc.save(filename)
    return filename


# =========================
# PDF (LIGHT VERSION)
# =========================

def generate_pdf(data):

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"Field_Visit_Report_{ts}.pdf"

    pdf = SimpleDocTemplate(filename)
    styles = getSampleStyleSheet()

    elements = [
        Paragraph("Field Visit Report", styles["Title"]),
        Spacer(1, 10),
        Paragraph("Summary version. Refer DOCX for full report.", styles["BodyText"])
    ]

    pdf.build(elements)
    return filename


# =========================
# DONE HANDLER
# =========================

async def done(update, context):

    try:
        docx_file = generate_docx(context.user_data)
        pdf_file = generate_pdf(context.user_data)

        await update.message.reply_text("📄 Generating Report...")

        with open(docx_file, "rb") as f:
            await update.message.reply_document(f)

        with open(pdf_file, "rb") as f:
            await update.message.reply_document(f)

        await update.message.reply_text("✅ Report Generated Successfully!")

    except Exception as e:
        await update.message.reply_text(f"❌ Error: {str(e)}")

    return ConversationHandler.END


# =========================
# CANCEL
# =========================

async def cancel(update, context):
    await update.message.reply_text("❌ Cancelled")
    return ConversationHandler.END


# =========================
# MAIN
# =========================

def main():

    app = Application.builder().token(TOKEN).build()

    conv = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
    LOCATION: [MessageHandler(filters.TEXT, location)],

    VISIT_DATE: [
        MessageHandler(filters.TEXT, visit_date)
    ],

    OBJECTIVE: [
        MessageHandler(filters.TEXT, objective)
    ],

    OBSERVATIONS: [
        MessageHandler(filters.TEXT, observations)
    ],

    ACTIONS: [
        MessageHandler(filters.TEXT, actions)
    ],

    PHOTOS: [
        MessageHandler(filters.PHOTO, photo_handler),
        CommandHandler("done", done),
    ],
},
        fallbacks=[CommandHandler("cancel", cancel)],
    )

    app.add_handler(conv)

    print("Field Visit Bot Running...")
    app.run_polling()


if __name__ == "__main__":
    main()